"""Export skill output to downloadable documents (DOCX, XLSX, CSV, PDF, TXT, MD)."""

import csv
import io
import re

from docx import Document
from docx.shared import Pt, RGBColor
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from fpdf import FPDF
from openpyxl import Workbook
from pydantic import BaseModel

from app.api.deps import verify_api_key

router = APIRouter()

SUPPORTED_FORMATS = {"docx", "xlsx", "csv", "pdf", "txt", "md"}


class ExportRequest(BaseModel):
    content: str
    filename: str = "output"
    format: str = "docx"


# ── DOCX ─────────────────────────────────────────────────────────────────────

def _to_docx(content: str) -> io.BytesIO:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)
        elif re.match(r"^[-*+] ", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, line[2:].strip())
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, re.sub(r"^\d+\. ", "", line).strip())
        elif line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph("\n".join(code_lines))
            p.runs[0].font.name = "Courier New"
            p.runs[0].font.size = Pt(9)
        elif _is_table_row(line):
            # Collect all table rows
            table_rows = []
            while i < len(lines) and (_is_table_row(lines[i]) or _is_table_sep(lines[i])):
                if not _is_table_sep(lines[i]):
                    cells = [c.strip() for c in lines[i].strip("|").split("|")]
                    table_rows.append(cells)
                i += 1
            if table_rows:
                max_cols = max(len(r) for r in table_rows)
                tbl = doc.add_table(rows=len(table_rows), cols=max_cols)
                tbl.style = "Table Grid"
                for r_idx, row in enumerate(table_rows):
                    for c_idx, cell in enumerate(row):
                        if c_idx < max_cols:
                            tbl.cell(r_idx, c_idx).text = cell
                            if r_idx == 0:
                                for para in tbl.cell(r_idx, c_idx).paragraphs:
                                    for run in para.runs:
                                        run.bold = True
            continue
        elif line.strip():
            p = doc.add_paragraph()
            _add_inline(p, line)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_inline(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")
    for part in pattern.split(text):
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4F)
        elif part:
            paragraph.add_run(part)


# ── XLSX ─────────────────────────────────────────────────────────────────────

def _to_xlsx(content: str) -> io.BytesIO:
    wb = Workbook()
    ws = wb.active
    assert ws is not None

    from openpyxl.styles import Font, PatternFill, Alignment
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="4F46E5")

    row_num = 1
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        if _is_table_row(line):
            # Write table rows
            is_first = True
            while i < len(lines) and (_is_table_row(lines[i]) or _is_table_sep(lines[i])):
                if not _is_table_sep(lines[i]):
                    cells = [c.strip() for c in lines[i].strip("|").split("|")]
                    for col_idx, cell in enumerate(cells, start=1):
                        cell_obj = ws.cell(row=row_num, column=col_idx, value=cell)
                        if is_first:
                            cell_obj.font = header_font
                            cell_obj.fill = header_fill
                            cell_obj.alignment = Alignment(horizontal="center")
                    row_num += 1
                    is_first = False
                i += 1
            row_num += 1  # blank row after table
            continue
        elif line.startswith("# ") or line.startswith("## ") or line.startswith("### "):
            text = re.sub(r"^#+\s+", "", line)
            cell_obj = ws.cell(row=row_num, column=1, value=text)
            cell_obj.font = Font(bold=True, size=12)
            row_num += 1
        elif line.strip():
            ws.cell(row=row_num, column=1, value=_strip_markdown(line))
            row_num += 1
        else:
            row_num += 1
        i += 1

    # Auto-size columns
    for col in ws.columns:  # type: ignore[union-attr]
        max_len = max((len(str(c.value or "")) for c in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 60)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ── CSV ──────────────────────────────────────────────────────────────────────

def _to_csv(content: str) -> io.BytesIO:
    buf = io.StringIO()
    writer = csv.writer(buf)
    for line in content.split("\n"):
        if _is_table_row(line):
            cells = [c.strip() for c in line.strip("|").split("|")]
            writer.writerow(cells)
        elif _is_table_sep(line):
            continue
        elif line.strip():
            writer.writerow([_strip_markdown(line)])
    return io.BytesIO(buf.getvalue().encode("utf-8-sig"))  # utf-8-sig for Excel compat


# ── PDF ──────────────────────────────────────────────────────────────────────

def _to_pdf(content: str, title: str) -> io.BytesIO:
    pdf = FPDF()
    pdf.set_margins(20, 20, 20)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    for line in content.split("\n"):
        plain = _strip_markdown(line)
        if not plain and not line.strip():
            pdf.ln(3)
            continue
        if line.startswith("# "):
            pdf.set_font("Helvetica", "B", 18)
        elif line.startswith("## "):
            pdf.set_font("Helvetica", "B", 14)
        elif line.startswith("### "):
            pdf.set_font("Helvetica", "B", 12)
        elif line.startswith("```"):
            continue
        elif re.match(r"^[-*+] ", line):
            pdf.set_font("Helvetica", "", 11)
            plain = f"  \u2022 {plain}"
        else:
            pdf.set_font("Helvetica", "", 11)

        if line.strip() in ("---", "***", "___"):
            pdf.ln(2)
        elif plain:
            pdf.multi_cell(0, 7, plain)
            pdf.ln(1)

    buf = io.BytesIO(pdf.output())
    buf.seek(0)
    return buf


# ── Helpers ───────────────────────────────────────────────────────────────────

def _is_table_row(line: str) -> bool:
    return "|" in line and line.strip().startswith("|")


def _is_table_sep(line: str) -> bool:
    return bool(re.match(r"^\s*\|[\s\-|:]+\|\s*$", line))


def _strip_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    text = re.sub(r"^#{1,6}\s+", "", text)
    text = re.sub(r"^[-*+] ", "", text)
    text = re.sub(r"^\d+\. ", "", text)
    return text.strip()


# ── Route ─────────────────────────────────────────────────────────────────────

MIME_TYPES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "csv": "text/csv",
    "pdf": "application/pdf",
    "txt": "text/plain; charset=utf-8",
    "md": "text/markdown; charset=utf-8",
}


@router.post("/export/{fmt}")
async def export_file(
    fmt: str,
    req: ExportRequest,
    _: None = Depends(verify_api_key),
) -> StreamingResponse:
    if fmt not in SUPPORTED_FORMATS:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {fmt}. Use one of: {', '.join(SUPPORTED_FORMATS)}")

    safe_name = re.sub(r"[^\w\- ]", "", req.filename).strip() or "output"

    if fmt == "docx":
        buf = _to_docx(req.content)
    elif fmt == "xlsx":
        buf = _to_xlsx(req.content)
    elif fmt == "csv":
        buf = _to_csv(req.content)
    elif fmt == "pdf":
        buf = _to_pdf(req.content, safe_name)
    elif fmt in ("txt", "md"):
        buf = io.BytesIO(req.content.encode("utf-8"))
    else:
        raise HTTPException(status_code=400, detail="Unsupported format")

    return StreamingResponse(
        buf,
        media_type=MIME_TYPES[fmt],
        headers={
            "Content-Disposition": f'attachment; filename="{safe_name}.{fmt}"',
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Expose-Headers": "Content-Disposition",
        },
    )
