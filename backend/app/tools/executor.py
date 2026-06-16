"""Tool execution handlers — called when Claude invokes a tool during skill execution."""

import io
import json
import logging
import subprocess
import sys
import textwrap
from pathlib import Path

logger = logging.getLogger(__name__)

OUTPUTS_DIR = Path(__file__).resolve().parents[3] / "outputs"


# ── Web Search ────────────────────────────────────────────────────────────────

def web_search(query: str, max_results: int = 5) -> str:
    try:
        from duckduckgo_search import DDGS
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=max_results))
        if not results:
            return "No results found."
        lines = []
        for i, r in enumerate(results, 1):
            lines.append(f"[{i}] {r['title']}\n{r['href']}\n{r['body']}\n")
        return "\n".join(lines)
    except Exception as e:
        return f"Search error: {e}"


# ── Read File ────────────────────────────────────────────────────────────────

def read_file(filename: str) -> str:
    safe = Path(filename).name
    path = OUTPUTS_DIR / safe
    if not path.exists():
        return f"File not found: {safe}"
    try:
        return path.read_text(encoding="utf-8")
    except Exception as e:
        return f"Read error: {e}"


# ── Run Code ─────────────────────────────────────────────────────────────────

def run_code(code: str, language: str = "python") -> str:
    if language.lower() != "python":
        return f"Only Python is supported for code execution (got: {language})"
    try:
        result = subprocess.run(
            [sys.executable, "-c", code],
            capture_output=True,
            text=True,
            timeout=15,
        )
        output = result.stdout.strip()
        errors = result.stderr.strip()
        parts = []
        if output:
            parts.append(f"stdout:\n{output}")
        if errors:
            parts.append(f"stderr:\n{errors}")
        if result.returncode != 0:
            parts.append(f"exit code: {result.returncode}")
        return "\n\n".join(parts) if parts else "(no output)"
    except subprocess.TimeoutExpired:
        return "Execution timed out (15s limit)"
    except Exception as e:
        return f"Execution error: {e}"


# ── Create Word Doc ───────────────────────────────────────────────────────────

def create_word_doc(filename: str, content: str, title: str = "") -> str:
    try:
        import docx
        from docx.shared import Pt

        safe = Path(filename).stem + ".docx"
        OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
        path = OUTPUTS_DIR / safe

        doc = docx.Document()

        if title:
            doc.add_heading(title, level=0)

        for block in content.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            # Headings
            if block.startswith("### "):
                doc.add_heading(block[4:], level=3)
            elif block.startswith("## "):
                doc.add_heading(block[3:], level=2)
            elif block.startswith("# "):
                doc.add_heading(block[2:], level=1)
            # Bullet list
            elif block.startswith("- ") or block.startswith("* "):
                for line in block.split("\n"):
                    text = line.lstrip("-* ").strip()
                    if text:
                        doc.add_paragraph(text, style="List Bullet")
            else:
                doc.add_paragraph(block)

        doc.save(str(path))
        return str(path)
    except Exception as e:
        return f"Word doc error: {e}"


# ── Create Excel File ─────────────────────────────────────────────────────────

def create_excel_file(filename: str, data: list[dict] | str, sheet_name: str = "Sheet1") -> str:
    """
    data can be:
      - list of dicts  → each dict is a row, keys become column headers
      - JSON string    → parsed to list of dicts
      - plain text     → each line is a row, comma-separated values
    """
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill
        from openpyxl.utils import get_column_letter

        safe = Path(filename).stem + ".xlsx"
        OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
        path = OUTPUTS_DIR / safe

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name

        # Parse data
        rows: list[list] = []
        headers: list[str] = []

        if isinstance(data, str):
            try:
                parsed = json.loads(data)
                if isinstance(parsed, list) and parsed and isinstance(parsed[0], dict):
                    data = parsed
            except json.JSONDecodeError:
                # plain CSV-like text
                lines = [l.strip() for l in data.strip().split("\n") if l.strip()]
                if lines:
                    headers = [h.strip() for h in lines[0].split(",")]
                    for line in lines[1:]:
                        rows.append([v.strip() for v in line.split(",")])

        if isinstance(data, list) and data and isinstance(data[0], dict):
            headers = list(data[0].keys())
            for item in data:
                rows.append([item.get(h, "") for h in headers])

        # Write headers
        if headers:
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill("solid", fgColor="4472C4")
                cell.font = Font(bold=True, color="FFFFFF")

        # Write rows
        for row_idx, row in enumerate(rows, 2 if headers else 1):
            for col_idx, value in enumerate(row, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)

        # Auto-width columns
        for col in ws.columns:
            max_len = max((len(str(c.value or "")) for c in col), default=10)
            ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 4, 50)

        wb.save(str(path))
        return str(path)
    except Exception as e:
        return f"Excel error: {e}"


# ── Dispatch ──────────────────────────────────────────────────────────────────

def execute_tool(name: str, inputs: dict) -> tuple[str, dict]:
    """
    Execute a tool by name. Returns (result_text, sse_event).
    sse_event is the event to stream back to the client.
    """
    if name == "web_search":
        query = inputs.get("query", "")
        max_results = int(inputs.get("max_results", 5))
        result = web_search(query, max_results)
        return result, {"type": "tool_used", "tool": "web_search", "query": query}

    elif name == "read_file":
        filename = inputs.get("filename", "")
        result = read_file(filename)
        return result, {"type": "tool_used", "tool": "read_file", "filename": filename}

    elif name == "run_code":
        code = inputs.get("code", "")
        language = inputs.get("language", "python")
        result = run_code(code, language)
        return result, {"type": "tool_used", "tool": "run_code", "language": language}

    elif name == "save_file":
        filename = inputs.get("filename", "output.txt")
        content = inputs.get("content", "")
        safe = Path(filename).name or "output.txt"
        OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
        file_path = OUTPUTS_DIR / safe
        file_path.write_text(content, encoding="utf-8")
        result = str(file_path)
        return result, {
            "type": "file_saved",
            "filename": safe,
            "path": result,
            "size": len(content),
        }

    elif name == "create_word_doc":
        filename = inputs.get("filename", "document.docx")
        content = inputs.get("content", "")
        title = inputs.get("title", "")
        path = create_word_doc(filename, content, title)
        return path, {
            "type": "file_saved",
            "filename": Path(filename).stem + ".docx",
            "path": path,
            "format": "docx",
        }

    elif name == "create_excel_file":
        filename = inputs.get("filename", "data.xlsx")
        data = inputs.get("data", [])
        sheet = inputs.get("sheet_name", "Sheet1")
        path = create_excel_file(filename, data, sheet)
        return path, {
            "type": "file_saved",
            "filename": Path(filename).stem + ".xlsx",
            "path": path,
            "format": "xlsx",
        }

    return f"Unknown tool: {name}", {"type": "error", "text": f"Unknown tool: {name}"}
